from __future__ import annotations

"""
ATS-ready resume/cover exporters: Markdown source → PDF + DOCX.
PDFs preserve heading/list/emphasis structure from the markdown.
DOCX uses plain sequential styles (ATS-friendly: no tables/textboxes).

Never silently succeed with markdown-only. Missing PDF/DOCX is a hard failure
after dependency install + retry.
"""

import logging
import re
import subprocess
import sys
import traceback
import zipfile
from pathlib import Path
from typing import Any

log = logging.getLogger("job_machine.export")

# Standard packet filenames (always written alongside markdown source)
RESUME_MD = "resume.md"
RESUME_PDF = "resume.pdf"
RESUME_DOCX = "resume.docx"
COVER_MD = "cover_letter.md"
COVER_PDF = "cover_letter.pdf"
COVER_DOCX = "cover_letter.docx"

STANDARD_PACKET_FILES = (
    RESUME_MD,
    RESUME_PDF,
    RESUME_DOCX,
    COVER_MD,
    COVER_PDF,
    COVER_DOCX,
)

REQUIRED_BINARY_FILES = (RESUME_PDF, RESUME_DOCX, COVER_PDF, COVER_DOCX)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HEADER_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^(\s*)([-*•]|\d+\.)\s+(.*)$")
_HR_RE = re.compile(r"^---+$")


class ExportGenerationError(RuntimeError):
    """Raised when PDF/DOCX generation fails after retries."""

    def __init__(self, message: str, *, errors: list[str] | None = None, folder: str | None = None):
        super().__init__(message)
        self.errors = errors or []
        self.folder = folder


def _strip_md_inline(text: str) -> str:
    t = _LINK_RE.sub(r"\1 (\2)", text)
    t = _BOLD_RE.sub(r"\1", t)
    t = _ITALIC_RE.sub(r"\1", t)
    return t.strip()


def _parse_blocks(markdown: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for raw in (markdown or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            blocks.append({"type": "blank"})
            continue
        if _HR_RE.match(line.strip()):
            blocks.append({"type": "hr"})
            continue
        hm = _HEADER_RE.match(line)
        if hm:
            level = len(hm.group(1))
            blocks.append({"type": "heading", "level": level, "text": _strip_md_inline(hm.group(2))})
            continue
        bm = _BULLET_RE.match(line)
        if bm:
            blocks.append(
                {
                    "type": "bullet",
                    "indent": len(bm.group(1) or ""),
                    "text": _strip_md_inline(bm.group(3)),
                }
            )
            continue
        blocks.append({"type": "para", "text": _strip_md_inline(line)})
    return blocks


def ensure_export_dependencies() -> dict[str, Any]:
    """Import reportlab + python-docx; pip install automatically if missing, then retry."""
    result: dict[str, Any] = {"reportlab": False, "python_docx": False, "installed": [], "errors": []}
    missing: list[str] = []

    try:
        import reportlab  # noqa: F401

        result["reportlab"] = True
        log.info("export.deps reportlab=ok version=%s", getattr(reportlab, "Version", "?"))
    except ImportError as exc:
        result["errors"].append(f"reportlab missing: {exc}")
        missing.append("reportlab>=4.2.0")
        log.warning("export.deps reportlab missing: %s", exc)

    try:
        import docx  # noqa: F401

        result["python_docx"] = True
        log.info("export.deps python-docx=ok")
    except ImportError as exc:
        result["errors"].append(f"python-docx missing: {exc}")
        missing.append("python-docx>=1.1.2")
        log.warning("export.deps python-docx missing: %s", exc)

    if not missing:
        return result

    log.warning("export.deps auto-installing: %s via %s", missing, sys.executable)
    try:
        cmd = [sys.executable, "-m", "pip", "install", *missing]
        proc = subprocess.run(cmd, capture_output=True, text=True, check=False)
        log.info("export.deps pip exit=%s stdout=%s stderr=%s", proc.returncode, proc.stdout[-2000:], proc.stderr[-2000:])
        if proc.returncode != 0:
            result["errors"].append(f"pip install failed ({proc.returncode}): {proc.stderr or proc.stdout}")
            return result
        result["installed"] = missing
    except Exception as exc:  # noqa: BLE001
        result["errors"].append(f"pip install exception: {exc}")
        log.exception("export.deps pip install failed")
        return result

    # Re-import after install
    try:
        import importlib

        import reportlab

        importlib.reload(reportlab)
        result["reportlab"] = True
    except Exception as exc:  # noqa: BLE001
        result["errors"].append(f"reportlab still unavailable after install: {exc}")
    try:
        import docx  # noqa: F401

        result["python_docx"] = True
    except Exception as exc:  # noqa: BLE001
        result["errors"].append(f"python-docx still unavailable after install: {exc}")

    return result


def write_markdown_pdf(path: Path, markdown: str, *, doc_title: str = "") -> None:
    """Render markdown to PDF. Raises on failure (never silent)."""
    if not (markdown or "").strip():
        raise ValueError("Cannot write PDF: markdown body is empty")

    deps = ensure_export_dependencies()
    if not deps.get("reportlab"):
        raise ImportError(
            "reportlab is required for PDF export. "
            + "; ".join(deps.get("errors") or ["install failed"])
        )

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        path.unlink()

    width, height = letter
    left, right = 54, width - 54
    max_width = right - left
    c = canvas.Canvas(str(path), pagesize=letter)
    if doc_title:
        c.setTitle(doc_title[:120])

    y = height - 54
    blocks = _parse_blocks(markdown)

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = height - 54

    def draw_wrapped(text: str, font: str, size: float, leading: float, indent: float = 0) -> None:
        nonlocal y
        c.setFont(font, size)
        words = (text or "").split()
        if not words:
            y -= leading
            return
        line = ""
        x = left + indent
        usable = max_width - indent
        for word in words:
            trial = f"{line} {word}".strip()
            if c.stringWidth(trial, font, size) <= usable:
                line = trial
            else:
                if y < 54:
                    new_page()
                    c.setFont(font, size)
                c.drawString(x, y, line)
                y -= leading
                line = word
        if line:
            if y < 54:
                new_page()
                c.setFont(font, size)
            c.drawString(x, y, line)
            y -= leading

    for block in blocks:
        kind = block["type"]
        if kind == "blank":
            y -= 8
            continue
        if kind == "hr":
            if y < 60:
                new_page()
            c.setStrokeColorRGB(0.55, 0.55, 0.55)
            c.setLineWidth(0.6)
            c.line(left, y + 4, right, y + 4)
            y -= 12
            continue
        if kind == "heading":
            level = int(block["level"])
            sizes = {1: 16, 2: 13, 3: 11}
            size = sizes.get(level, 11)
            y -= 6 if level == 1 else 4
            if y < 60:
                new_page()
            draw_wrapped(block["text"], "Helvetica-Bold", size, size + 4)
            y -= 2
            continue
        if kind == "bullet":
            indent = 14 + min(block.get("indent", 0), 24)
            draw_wrapped(f"• {block['text']}", "Helvetica", 10, 13, indent=indent - 14)
            continue
        draw_wrapped(block["text"], "Helvetica", 10, 13)

    c.save()
    if not path.exists() or path.stat().st_size < 100:
        raise RuntimeError(f"PDF write produced empty/missing file: {path}")
    log.info("export.pdf wrote %s (%s bytes)", path, path.stat().st_size)


def write_markdown_docx(path: Path, markdown: str, *, doc_title: str = "") -> None:
    """ATS-friendly DOCX. Raises on failure (never silent)."""
    if not (markdown or "").strip():
        raise ValueError("Cannot write DOCX: markdown body is empty")

    deps = ensure_export_dependencies()
    if not deps.get("python_docx"):
        raise ImportError(
            "python-docx is required for DOCX export. "
            + "; ".join(deps.get("errors") or ["install failed"])
        )

    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        path.unlink()

    doc = Document()
    if doc_title:
        doc.core_properties.title = doc_title[:200]

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for block in _parse_blocks(markdown):
        kind = block["type"]
        if kind == "blank":
            doc.add_paragraph("")
            continue
        if kind == "hr":
            doc.add_paragraph("—" * 24)
            continue
        if kind == "heading":
            level = min(int(block["level"]), 3)
            p = doc.add_heading(block["text"], level=level)
            for run in p.runs:
                run.font.name = "Calibri"
            continue
        if kind == "bullet":
            p = doc.add_paragraph(block["text"], style="List Bullet")
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            continue
        p = doc.add_paragraph(block["text"])
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    doc.save(str(path))
    if not path.exists() or path.stat().st_size < 100:
        raise RuntimeError(f"DOCX write produced empty/missing file: {path}")
    log.info("export.docx wrote %s (%s bytes)", path, path.stat().st_size)


def _try_write(
    label: str,
    fn,
    path: Path,
    markdown: str,
    doc_title: str,
    errors: list[str],
) -> bool:
    try:
        log.info("export.generate start %s -> %s", label, path)
        fn(path, markdown, doc_title=doc_title)
        ok = path.exists() and path.stat().st_size > 0
        log.info("export.generate ok %s exists=%s size=%s", label, ok, path.stat().st_size if path.exists() else 0)
        return ok
    except Exception as exc:  # noqa: BLE001
        detail = f"{label}: {type(exc).__name__}: {exc}"
        errors.append(detail)
        log.error("export.generate FAIL %s\n%s", detail, traceback.format_exc())
        return False


def write_ats_packet_documents(
    folder: Path,
    *,
    resume_markdown: str,
    cover_markdown: str,
    company: str = "",
    title: str = "",
    require_all: bool = True,
) -> dict[str, Any]:
    """
    Write all six packet files into folder.
    If require_all=True (default), raises ExportGenerationError unless PDF+DOCX exist.
    """
    folder = Path(folder).resolve()
    folder.mkdir(parents=True, exist_ok=True)
    log.info(
        "export.packet start folder=%s company=%s title=%s resume_chars=%s cover_chars=%s",
        folder,
        company,
        title,
        len(resume_markdown or ""),
        len(cover_markdown or ""),
    )

    deps = ensure_export_dependencies()
    log.info("export.deps status=%s", deps)

    resume_md_path = folder / RESUME_MD
    cover_md_path = folder / COVER_MD
    resume_md_path.write_text(resume_markdown or "", encoding="utf-8")
    cover_md_path.write_text(cover_markdown or "", encoding="utf-8")
    log.info("export.md wrote %s (%s bytes) %s (%s bytes)", resume_md_path, resume_md_path.stat().st_size, cover_md_path, cover_md_path.stat().st_size)

    resume_label = f"Resume — {title} @ {company}".strip(" —@")
    cover_label = f"Cover Letter — {title} @ {company}".strip(" —@")

    resume_pdf = folder / RESUME_PDF
    cover_pdf = folder / COVER_PDF
    resume_docx = folder / RESUME_DOCX
    cover_docx = folder / COVER_DOCX

    errors: list[str] = list(deps.get("errors") or [])

    def _generate_all() -> tuple[bool, bool, bool, bool]:
        r_pdf = _try_write("resume.pdf", write_markdown_pdf, resume_pdf, resume_markdown, resume_label, errors)
        c_pdf = _try_write("cover_letter.pdf", write_markdown_pdf, cover_pdf, cover_markdown, cover_label, errors)
        r_docx = _try_write("resume.docx", write_markdown_docx, resume_docx, resume_markdown, resume_label, errors)
        c_docx = _try_write("cover_letter.docx", write_markdown_docx, cover_docx, cover_markdown, cover_label, errors)
        return r_pdf, c_pdf, r_docx, c_docx

    resume_pdf_ok, cover_pdf_ok, resume_docx_ok, cover_docx_ok = _generate_all()

    # If any binary failed, re-ensure deps and retry once
    if not (resume_pdf_ok and cover_pdf_ok and resume_docx_ok and cover_docx_ok):
        log.warning("export.packet incomplete — installing deps and retrying once. errors=%s", errors)
        ensure_export_dependencies()
        resume_pdf_ok, cover_pdf_ok, resume_docx_ok, cover_docx_ok = _generate_all()

    files_written = [n for n in STANDARD_PACKET_FILES if (folder / n).exists() and (folder / n).stat().st_size > 0]
    missing = [n for n in STANDARD_PACKET_FILES if n not in files_written]
    verified = {
        n: {
            "path": str(folder / n),
            "exists": (folder / n).exists(),
            "bytes": (folder / n).stat().st_size if (folder / n).exists() else 0,
        }
        for n in STANDARD_PACKET_FILES
    }

    preferred_resume = RESUME_PDF if resume_pdf_ok else (RESUME_DOCX if resume_docx_ok else RESUME_MD)
    preferred_cover = COVER_PDF if cover_pdf_ok else (COVER_DOCX if cover_docx_ok else COVER_MD)

    result = {
        "folder": str(folder),
        "absolute_folder": str(folder),
        "paths": {
            "resume_md": str(resume_md_path),
            "cover_md": str(cover_md_path),
            "resume_pdf": str(resume_pdf) if resume_pdf_ok else None,
            "cover_pdf": str(cover_pdf) if cover_pdf_ok else None,
            "resume_docx": str(resume_docx) if resume_docx_ok else None,
            "cover_docx": str(cover_docx) if cover_docx_ok else None,
        },
        "files_written": files_written,
        "missing_files": missing,
        "verified_on_disk": verified,
        "preferred_resume": preferred_resume,
        "preferred_cover": preferred_cover,
        "preferred_resume_path": str(folder / preferred_resume),
        "preferred_cover_path": str(folder / preferred_cover),
        "pdf_ok": resume_pdf_ok and cover_pdf_ok,
        "docx_ok": resume_docx_ok and cover_docx_ok,
        "fallback_used": False,
        "errors": errors,
        "deps": deps,
        "success": len(missing) == 0,
    }

    log.info(
        "export.packet done folder=%s success=%s files_written=%s missing=%s errors=%s",
        folder,
        result["success"],
        files_written,
        missing,
        errors,
    )

    if require_all and missing:
        msg = (
            f"Export incomplete — missing {missing} in {folder}. "
            f"Errors: {errors or ['unknown']}"
        )
        raise ExportGenerationError(msg, errors=errors, folder=str(folder))

    return result


def build_packet_zip(folder: Path, zip_path: Path | None = None) -> Path:
    """Zip ATS packet files for one-click download. Requires all six files when present."""
    folder = Path(folder)
    zip_path = zip_path or (folder / "export_packet.zip")
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name in STANDARD_PACKET_FILES:
            p = folder / name
            if p.exists():
                zf.write(p, arcname=name)
                log.info("export.zip add %s (%s bytes)", name, p.stat().st_size)
        meta = folder / "meta.json"
        if meta.exists():
            zf.write(meta, arcname="meta.json")
    log.info("export.zip wrote %s (%s bytes)", zip_path, zip_path.stat().st_size)
    return zip_path


def is_standard_packet_filename(name: str) -> bool:
    n = (name or "").lower().strip()
    return n in {x.lower() for x in STANDARD_PACKET_FILES} or n in {
        "export_packet.zip",
        "cover.md",
        "cover.pdf",
        "cover.docx",
    }
